#!/usr/bin/env python3
"""
Substack Source Linker
======================

Purpose
-------
Take one Markdown article and one CSV source map from an article folder,
insert first-use hyperlinks, then create publication-ready files for Substack.

Folder structure
----------------

Substack_Source_Linker/
│
├── substack_source_linker.py
├── Originals/
│   └── Glitter_Conspiracy/
│       ├── ANY_FILENAME.md
│       └── ANY_FILENAME.csv
│
└── Completed/
    └── Glitter_Conspiracy/
        ├── <article_name>_substack.docx
        ├── <article_name>_substack.html
        └── <article_name>_link_report.txt

Usage
-----

    python substack_source_linker.py Glitter_Conspiracy

The filenames inside the article directory do NOT matter.

Requirements
------------

    pip install python-docx

Optional, for higher-fidelity HTML conversion:

    pip install markdown

The program will:
- automatically find exactly one .md/.markdown file
- automatically find exactly one .csv file
- insert each source link only at the first eligible occurrence
- preserve the original Markdown source file
- create a Word document designed for copy/paste into Substack
- create an HTML backup/preview
- create a link report showing any unmatched source phrases
"""

from __future__ import annotations

import argparse
import csv
import html
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple
from urllib.parse import parse_qsl, urlencode, urlsplit, urlunsplit


ROOT = Path(__file__).resolve().parent
ORIGINALS_DIR = ROOT / "Originals"
COMPLETED_DIR = ROOT / "Completed"


# ---------------------------------------------------------------------------
# Dependencies
# ---------------------------------------------------------------------------

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print()
    print("ERROR: python-docx is required to create the Substack Word file.")
    print()
    print("Install it with:")
    print()
    print("    pip install python-docx")
    print()
    sys.exit(1)


# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------

@dataclass
class Source:
    number: str
    phrase: str
    link: str


@dataclass
class InlineSegment:
    text: str
    bold: bool = False
    italic: bool = False
    url: Optional[str] = None


# ---------------------------------------------------------------------------
# File discovery
# ---------------------------------------------------------------------------

def find_exactly_one_file(
    folder: Path,
    extensions: set[str],
    label: str,
) -> Path:
    matches = sorted(
        path
        for path in folder.iterdir()
        if path.is_file() and path.suffix.lower() in extensions
    )

    if not matches:
        expected = ", ".join(sorted(extensions))
        raise FileNotFoundError(
            f"No {label} found in:\n"
            f"  {folder}\n\n"
            f"Expected exactly one file ending in: {expected}"
        )

    if len(matches) > 1:
        names = "\n".join(f"  - {path.name}" for path in matches)
        raise RuntimeError(
            f"Found more than one {label} in:\n"
            f"  {folder}\n\n"
            f"{names}\n\n"
            f"The folder must contain exactly one {label}."
        )

    return matches[0]


# ---------------------------------------------------------------------------
# Source CSV
# ---------------------------------------------------------------------------

def normalize_url(url: str) -> str:
    """Remove common tracking parameters while preserving useful parameters."""
    url = url.strip()

    if not url:
        return url

    try:
        parts = urlsplit(url)
    except Exception:
        return url

    tracking_keys = {
        "source",
        "campaign",
        "ref",
        "ref_src",
        "fbclid",
        "gclid",
        "mc_cid",
        "mc_eid",
    }

    kept = []

    for key, value in parse_qsl(parts.query, keep_blank_values=True):
        key_lower = key.lower()

        if key_lower.startswith("utm_"):
            continue

        if key_lower in tracking_keys:
            continue

        kept.append((key, value))

    return urlunsplit(
        (
            parts.scheme,
            parts.netloc,
            parts.path,
            urlencode(kept, doseq=True),
            parts.fragment,
        )
    )


def read_sources(csv_path: Path) -> List[Source]:
    with csv_path.open(
        "r",
        encoding="utf-8-sig",
        newline="",
    ) as handle:
        rows = [
            row
            for row in csv.reader(handle)
            if row and any(cell.strip() for cell in row)
        ]

    if not rows:
        raise ValueError(f"The CSV is empty: {csv_path}")

    possible_headers = {
        "source",
        "source_number",
        "number",
        "#",
        "phrase",
        "name",
        "anchor",
        "link",
        "url",
    }

    first_row = [cell.strip().lower() for cell in rows[0]]
    has_header = any(cell in possible_headers for cell in first_row)
    data_rows = rows[1:] if has_header else rows

    sources: List[Source] = []

    for row_num, row in enumerate(
        data_rows,
        start=2 if has_header else 1,
    ):
        if len(row) < 3:
            print(
                f"WARNING: skipping CSV row {row_num}; "
                f"expected at least 3 columns.",
                file=sys.stderr,
            )
            continue

        number = row[0].strip()
        phrase = row[1].strip()
        link = normalize_url(row[2].strip())

        if not phrase or not link:
            print(
                f"WARNING: skipping CSV row {row_num}; "
                f"phrase or link is blank.",
                file=sys.stderr,
            )
            continue

        sources.append(
            Source(
                number=number,
                phrase=phrase,
                link=link,
            )
        )

    if not sources:
        raise ValueError(
            f"No usable source rows were found in {csv_path}"
        )

    return sources


# ---------------------------------------------------------------------------
# First-use link insertion
# ---------------------------------------------------------------------------

def is_inside_markdown_link(
    text: str,
    start: int,
    end: int,
) -> bool:
    line_start = text.rfind("\n", 0, start) + 1
    line_end = text.find("\n", end)

    if line_end == -1:
        line_end = len(text)

    segment = text[line_start:line_end]
    local_start = start - line_start
    local_end = end - line_start

    for match in re.finditer(
        r"\[[^\]]+\]\([^)]+\)",
        segment,
    ):
        if (
            match.start() <= local_start
            and local_end <= match.end()
        ):
            return True

    return False


def is_inside_fenced_code(
    text: str,
    position: int,
) -> bool:
    return text[:position].count("```") % 2 == 1


def replace_first_unlinked(
    text: str,
    phrase: str,
    url: str,
) -> Tuple[str, bool]:
    if not phrase:
        return text, False

    left_boundary = r"(?<!\w)" if phrase[0].isalnum() else ""
    right_boundary = r"(?!\w)" if phrase[-1].isalnum() else ""

    pattern = re.compile(
        left_boundary + re.escape(phrase) + right_boundary,
        re.IGNORECASE,
    )

    for match in pattern.finditer(text):
        if is_inside_markdown_link(
            text,
            match.start(),
            match.end(),
        ):
            continue

        if is_inside_fenced_code(
            text,
            match.start(),
        ):
            continue

        original = match.group(0)
        replacement = f"[{original}]({url})"

        return (
            text[:match.start()]
            + replacement
            + text[match.end():],
            True,
        )

    return text, False


def insert_links(
    article_text: str,
    sources: List[Source],
):
    # Longer phrases first prevents shorter anchors from consuming
    # text intended for a longer source phrase.
    ordered = sorted(
        enumerate(sources),
        key=lambda item: (
            -len(item[1].phrase),
            item[0],
        ),
    )

    results = {}
    text = article_text

    for original_index, source in ordered:
        text, success = replace_first_unlinked(
            text,
            source.phrase,
            source.link,
        )
        results[original_index] = success

    linked = []
    missing = []

    for index, source in enumerate(sources):
        if results.get(index):
            linked.append(source)
        else:
            missing.append(source)

    return text, linked, missing


# ---------------------------------------------------------------------------
# Inline Markdown parsing
# ---------------------------------------------------------------------------

INLINE_PATTERN = re.compile(
    r"""
    \[([^\]]+)\]\((https?://[^)]+)\)     # markdown link
    |
    \*\*(.+?)\*\*                        # bold
    |
    (?<!\*)\*([^*\n]+?)\*(?!\*)          # italic
    """,
    re.VERBOSE,
)


def parse_inline(
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    url: Optional[str] = None,
) -> List[InlineSegment]:
    """
    Convert common inline Markdown into styled segments.

    Handles combinations such as:
        **bold**
        *italic*
        [link](https://...)
        **[bold link](https://...)**
        [**bold link**](https://...)
    """
    segments: List[InlineSegment] = []
    cursor = 0

    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            segments.append(
                InlineSegment(
                    text=text[cursor:match.start()],
                    bold=bold,
                    italic=italic,
                    url=url,
                )
            )

        link_label = match.group(1)
        link_url = match.group(2)
        bold_inner = match.group(3)
        italic_inner = match.group(4)

        if link_label is not None:
            segments.extend(
                parse_inline(
                    link_label,
                    bold=bold,
                    italic=italic,
                    url=link_url,
                )
            )

        elif bold_inner is not None:
            segments.extend(
                parse_inline(
                    bold_inner,
                    bold=True,
                    italic=italic,
                    url=url,
                )
            )

        elif italic_inner is not None:
            segments.extend(
                parse_inline(
                    italic_inner,
                    bold=bold,
                    italic=True,
                    url=url,
                )
            )

        cursor = match.end()

    if cursor < len(text):
        segments.append(
            InlineSegment(
                text=text[cursor:],
                bold=bold,
                italic=italic,
                url=url,
            )
        )

    return segments


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def add_hyperlink_run(
    paragraph,
    text: str,
    url: str,
    *,
    bold: bool = False,
    italic: bool = False,
):
    """
    Add a real clickable hyperlink to a python-docx paragraph.
    """
    part = paragraph.part

    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    # Blue + underline gives a normal visual hyperlink in Word.
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    if bold:
        bold_element = OxmlElement("w:b")
        run_properties.append(bold_element)

    if italic:
        italic_element = OxmlElement("w:i")
        run_properties.append(italic_element)

    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run.append(text_element)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_to_paragraph(
    paragraph,
    text: str,
):
    segments = parse_inline(text)

    for segment in segments:
        if not segment.text:
            continue

        if segment.url:
            add_hyperlink_run(
                paragraph,
                segment.text,
                segment.url,
                bold=segment.bold,
                italic=segment.italic,
            )
        else:
            run = paragraph.add_run(segment.text)
            run.bold = segment.bold
            run.italic = segment.italic


def clean_heading_markup(text: str) -> str:
    """
    Heading styles already provide emphasis, so remove outer ** markers
    while preserving any links and inline formatting inside.
    """
    stripped = text.strip()

    if stripped.startswith("**") and stripped.endswith("**"):
        return stripped[2:-2]

    return stripped


def set_document_defaults(document: Document):
    section = document.sections[0]

    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11.5)

    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.12

    for level in range(1, 4):
        style_name = f"Heading {level}"

        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "Arial"
            style.font.bold = True

            if level == 1:
                style.font.size = Pt(20)
            elif level == 2:
                style.font.size = Pt(15)
            else:
                style.font.size = Pt(13)


def create_quote_paragraph(
    document: Document,
    content: str,
):
    try:
        paragraph = document.add_paragraph(style="Quote")
    except KeyError:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.35)

    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)

    add_inline_to_paragraph(
        paragraph,
        content,
    )

    return paragraph


def create_image_placeholder(
    document: Document,
    text: str,
):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)

    run = paragraph.add_run(text)
    run.bold = True
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    return paragraph


def split_markdown_table_row(line: str) -> List[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def is_markdown_table_separator(line: str) -> bool:
    cells = split_markdown_table_row(line)

    if not cells:
        return False

    return all(
        bool(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")))
        for cell in cells
    )


def add_markdown_table(
    document: Document,
    table_lines: List[str],
):
    rows = [
        split_markdown_table_row(line)
        for line in table_lines
        if not is_markdown_table_separator(line)
    ]

    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(
        rows=len(rows),
        cols=column_count,
    )
    table.style = "Table Grid"

    for row_index, cells in enumerate(rows):
        for col_index in range(column_count):
            value = cells[col_index] if col_index < len(cells) else ""
            cell = table.cell(row_index, col_index)
            cell.text = ""

            paragraph = cell.paragraphs[0]
            segments = parse_inline(value)

            for segment in segments:
                if segment.url:
                    add_hyperlink_run(
                        paragraph,
                        segment.text,
                        segment.url,
                        bold=segment.bold or row_index == 0,
                        italic=segment.italic,
                    )
                else:
                    run = paragraph.add_run(segment.text)
                    run.bold = segment.bold or row_index == 0
                    run.italic = segment.italic

    document.add_paragraph()


def markdown_to_docx(
    md_text: str,
    output_path: Path,
):
    document = Document()
    set_document_defaults(document)

    lines = md_text.splitlines()
    index = 0
    paragraph_buffer: List[str] = []

    def flush_paragraph_buffer():
        nonlocal paragraph_buffer

        if not paragraph_buffer:
            return

        text = " ".join(
            line.strip()
            for line in paragraph_buffer
        ).strip()

        if text:
            paragraph = document.add_paragraph()
            add_inline_to_paragraph(
                paragraph,
                text,
            )

        paragraph_buffer = []

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()
        stripped = line.strip()

        # Blank line
        if not stripped:
            flush_paragraph_buffer()
            index += 1
            continue

        # Horizontal rule
        if re.fullmatch(r"\s*---+\s*", line):
            flush_paragraph_buffer()
            index += 1
            continue

        # Markdown table
        if (
            "|" in line
            and index + 1 < len(lines)
            and is_markdown_table_separator(lines[index + 1])
        ):
            flush_paragraph_buffer()

            table_lines = [line, lines[index + 1]]
            index += 2

            while (
                index < len(lines)
                and "|" in lines[index]
                and lines[index].strip()
            ):
                table_lines.append(lines[index])
                index += 1

            add_markdown_table(
                document,
                table_lines,
            )
            continue

        # Heading
        heading_match = re.match(
            r"^(#{1,6})\s+(.+)$",
            line,
        )

        if heading_match:
            flush_paragraph_buffer()

            level = min(
                len(heading_match.group(1)),
                3,
            )

            heading_text = clean_heading_markup(
                heading_match.group(2)
            )

            paragraph = document.add_paragraph(
                style=f"Heading {level}"
            )

            add_inline_to_paragraph(
                paragraph,
                heading_text,
            )

            index += 1
            continue

        # Block quote
        quote_match = re.match(
            r"^>\s?(.*)$",
            line,
        )

        if quote_match:
            flush_paragraph_buffer()

            quote_parts = [
                quote_match.group(1).strip()
            ]

            index += 1

            while index < len(lines):
                next_quote = re.match(
                    r"^>\s?(.*)$",
                    lines[index],
                )

                if not next_quote:
                    break

                quote_parts.append(
                    next_quote.group(1).strip()
                )
                index += 1

            create_quote_paragraph(
                document,
                " ".join(quote_parts),
            )
            continue

        # Image/editorial placeholder
        if (
            stripped.startswith("[IMAGE:")
            and stripped.endswith("]")
        ):
            flush_paragraph_buffer()
            create_image_placeholder(
                document,
                stripped,
            )
            index += 1
            continue

        # Unordered list
        unordered_match = re.match(
            r"^\s*[-*+]\s+(.+)$",
            line,
        )

        if unordered_match:
            flush_paragraph_buffer()

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            add_inline_to_paragraph(
                paragraph,
                unordered_match.group(1),
            )

            index += 1
            continue

        # Ordered list
        ordered_match = re.match(
            r"^\s*\d+[.)]\s+(.+)$",
            line,
        )

        if ordered_match:
            flush_paragraph_buffer()

            paragraph = document.add_paragraph(
                style="List Number"
            )

            add_inline_to_paragraph(
                paragraph,
                ordered_match.group(1),
            )

            index += 1
            continue

        paragraph_buffer.append(line)
        index += 1

    flush_paragraph_buffer()

    document.save(output_path)


# ---------------------------------------------------------------------------
# HTML conversion
# ---------------------------------------------------------------------------

def markdown_to_html(md_text: str) -> str:
    """
    Prefer Python-Markdown if installed. Its 'extra' extension handles
    tables, lists, blockquotes, headings, links, bold, and italics well.
    """
    try:
        import markdown

        return markdown.markdown(
            md_text,
            extensions=[
                "extra",
                "sane_lists",
                "tables",
            ],
        )
    except ImportError:
        pass

    # Lightweight fallback. The DOCX remains the primary output.
    lines = md_text.splitlines()
    output: List[str] = []
    paragraph_buffer: List[str] = []
    index = 0
    in_ul = False
    in_ol = False

    def inline(text: str) -> str:
        segments = parse_inline(text)
        result = []

        for segment in segments:
            content = html.escape(
                segment.text,
                quote=False,
            )

            if segment.bold:
                content = f"<strong>{content}</strong>"

            if segment.italic:
                content = f"<em>{content}</em>"

            if segment.url:
                safe_url = html.escape(
                    segment.url,
                    quote=True,
                )
                content = (
                    f'<a href="{safe_url}">'
                    f"{content}"
                    f"</a>"
                )

            result.append(content)

        return "".join(result)

    def flush_paragraph():
        nonlocal paragraph_buffer

        if paragraph_buffer:
            combined = " ".join(
                line.strip()
                for line in paragraph_buffer
            ).strip()

            if combined:
                output.append(
                    f"<p>{inline(combined)}</p>"
                )

            paragraph_buffer = []

    def close_lists():
        nonlocal in_ul, in_ol

        if in_ul:
            output.append("</ul>")
            in_ul = False

        if in_ol:
            output.append("</ol>")
            in_ol = False

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            close_lists()
            index += 1
            continue

        if re.fullmatch(r"\s*---+\s*", line):
            flush_paragraph()
            close_lists()
            output.append("<hr>")
            index += 1
            continue

        # Table
        if (
            "|" in line
            and index + 1 < len(lines)
            and is_markdown_table_separator(lines[index + 1])
        ):
            flush_paragraph()
            close_lists()

            header = split_markdown_table_row(line)
            index += 2

            rows = []

            while (
                index < len(lines)
                and "|" in lines[index]
                and lines[index].strip()
            ):
                rows.append(
                    split_markdown_table_row(
                        lines[index]
                    )
                )
                index += 1

            output.append("<table>")
            output.append("<thead><tr>")

            for cell in header:
                output.append(
                    f"<th>{inline(cell)}</th>"
                )

            output.append("</tr></thead>")
            output.append("<tbody>")

            for row in rows:
                output.append("<tr>")

                for cell in row:
                    output.append(
                        f"<td>{inline(cell)}</td>"
                    )

                output.append("</tr>")

            output.append("</tbody></table>")
            continue

        heading = re.match(
            r"^(#{1,6})\s+(.+)$",
            line,
        )

        if heading:
            flush_paragraph()
            close_lists()

            level = len(
                heading.group(1)
            )

            heading_text = clean_heading_markup(
                heading.group(2)
            )

            output.append(
                f"<h{level}>"
                f"{inline(heading_text)}"
                f"</h{level}>"
            )

            index += 1
            continue

        quote = re.match(
            r"^>\s?(.*)$",
            line,
        )

        if quote:
            flush_paragraph()
            close_lists()

            quote_parts = [
                quote.group(1).strip()
            ]

            index += 1

            while index < len(lines):
                next_quote = re.match(
                    r"^>\s?(.*)$",
                    lines[index],
                )

                if not next_quote:
                    break

                quote_parts.append(
                    next_quote.group(1).strip()
                )
                index += 1

            output.append(
                "<blockquote>"
                + inline(" ".join(quote_parts))
                + "</blockquote>"
            )

            continue

        if (
            stripped.startswith("[IMAGE:")
            and stripped.endswith("]")
        ):
            flush_paragraph()
            close_lists()

            output.append(
                '<p class="image-placeholder">'
                + html.escape(stripped)
                + "</p>"
            )

            index += 1
            continue

        unordered = re.match(
            r"^\s*[-*+]\s+(.+)$",
            line,
        )

        if unordered:
            flush_paragraph()

            if in_ol:
                output.append("</ol>")
                in_ol = False

            if not in_ul:
                output.append("<ul>")
                in_ul = True

            output.append(
                f"<li>{inline(unordered.group(1))}</li>"
            )

            index += 1
            continue

        ordered = re.match(
            r"^\s*\d+[.)]\s+(.+)$",
            line,
        )

        if ordered:
            flush_paragraph()

            if in_ul:
                output.append("</ul>")
                in_ul = False

            if not in_ol:
                output.append("<ol>")
                in_ol = True

            output.append(
                f"<li>{inline(ordered.group(1))}</li>"
            )

            index += 1
            continue

        paragraph_buffer.append(line)
        index += 1

    flush_paragraph()
    close_lists()

    return "\n".join(output)


def wrap_html(
    body: str,
    title: str,
) -> str:
    return f"""<!doctype html>
<html lang="en">

<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">

<title>{html.escape(title)}</title>

<style>
body {{
    font-family: Georgia, "Times New Roman", serif;
    font-size: 20px;
    line-height: 1.65;
    max-width: 780px;
    margin: 48px auto;
    padding: 0 24px;
    color: #171717;
}}

h1,
h2,
h3,
h4,
h5,
h6 {{
    font-family: Arial, Helvetica, sans-serif;
    line-height: 1.18;
    margin-top: 1.7em;
}}

h1 {{
    font-size: 2.1em;
}}

h2 {{
    font-size: 1.55em;
}}

h3 {{
    font-size: 1.3em;
}}

p {{
    margin: 1em 0;
}}

a {{
    color: inherit;
    text-decoration: underline;
}}

blockquote {{
    margin: 1.4em 0;
    padding: 0.1em 0 0.1em 1.1em;
    border-left: 3px solid #888;
    font-style: italic;
}}

li {{
    margin: 0.35em 0;
}}

table {{
    border-collapse: collapse;
    width: 100%;
    margin: 1.3em 0;
}}

th,
td {{
    border: 1px solid #bbb;
    padding: 0.55em;
    text-align: left;
    vertical-align: top;
}}

th {{
    font-weight: bold;
}}

.image-placeholder {{
    font-family: Arial, Helvetica, sans-serif;
    font-size: 0.85em;
    font-weight: bold;
    font-style: italic;
    color: #666;
    background: #f3f3f3;
    padding: 0.7em;
}}
</style>

</head>

<body>

{body}

</body>
</html>
"""


# ---------------------------------------------------------------------------
# Report
# ---------------------------------------------------------------------------

def create_report(
    project_name: str,
    article_file: Path,
    csv_file: Path,
    linked: List[Source],
    missing: List[Source],
) -> str:
    lines = [
        f"Project: {project_name}",
        f"Original article: {article_file.name}",
        f"Source CSV: {csv_file.name}",
        "",
        f"Sources loaded: {len(linked) + len(missing)}",
        f"Links inserted: {len(linked)}",
        f"Sources not matched: {len(missing)}",
        "",
    ]

    if linked:
        lines.append("SUCCESSFULLY LINKED:")

        for source in linked:
            lines.append(
                f"{source.number}\t"
                f"{source.phrase}\t"
                f"{source.link}"
            )

        lines.append("")

    if missing:
        lines.append("NOT FOUND / NOT INSERTED:")

        for source in missing:
            lines.append(
                f"{source.number}\t"
                f"{source.phrase}\t"
                f"{source.link}"
            )
    else:
        lines.append(
            "All source phrases were found and linked."
        )

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description=(
            "Process the single Markdown article and CSV source map "
            "inside an Originals article folder, then create a DOCX "
            "and HTML version for Substack."
        )
    )

    parser.add_argument(
        "directory_name",
        help=(
            "Folder name inside Originals. "
            "Example: Glitter_Conspiracy"
        ),
    )

    args = parser.parse_args()

    project_name = args.directory_name.strip()
    source_folder = ORIGINALS_DIR / project_name
    completed_folder = COMPLETED_DIR / project_name

    if not source_folder.exists():
        print()
        print("ERROR: Article folder does not exist.")
        print()
        print("Expected:")
        print(f"  {source_folder}")
        print()
        sys.exit(1)

    if not source_folder.is_dir():
        print(
            f"ERROR: Not a directory: "
            f"{source_folder}"
        )
        sys.exit(1)

    try:
        article_file = find_exactly_one_file(
            source_folder,
            {".md", ".markdown"},
            "Markdown file",
        )

        csv_file = find_exactly_one_file(
            source_folder,
            {".csv"},
            "CSV file",
        )

        print()
        print(f"Found article: {article_file.name}")
        print(f"Found sources: {csv_file.name}")
        print()

        article_text = article_file.read_text(
            encoding="utf-8-sig"
        )

        sources = read_sources(csv_file)

        linked_markdown, linked, missing = insert_links(
            article_text,
            sources,
        )

        completed_folder.mkdir(
            parents=True,
            exist_ok=True,
        )

        stem = article_file.stem

        docx_output = (
            completed_folder
            / f"{stem}_substack.docx"
        )

        html_output = (
            completed_folder
            / f"{stem}_substack.html"
        )

        report_output = (
            completed_folder
            / f"{stem}_link_report.txt"
        )

        # Primary output: Word / rich text
        markdown_to_docx(
            linked_markdown,
            docx_output,
        )

        # Backup / browser-copy output
        html_body = markdown_to_html(
            linked_markdown
        )

        html_page = wrap_html(
            html_body,
            stem.replace("_", " ")
                .replace("-", " ")
                .title(),
        )

        html_output.write_text(
            html_page,
            encoding="utf-8",
        )

        # Link audit
        report = create_report(
            project_name,
            article_file,
            csv_file,
            linked,
            missing,
        )

        report_output.write_text(
            report,
            encoding="utf-8",
        )

    except Exception as exc:
        print()
        print(f"ERROR: {exc}")
        print()
        sys.exit(1)

    print("=" * 70)
    print("SUBSTACK SOURCE LINKER — COMPLETE")
    print("=" * 70)
    print()

    print(f"Project folder: {project_name}")
    print(f"Input Markdown: {article_file.name}")
    print(f"Input CSV:      {csv_file.name}")
    print()

    print(f"Links inserted: {len(linked)}")
    print(f"Not matched:    {len(missing)}")
    print()

    print("Created:")
    print(f"  Word:   {docx_output}")
    print(f"  HTML:   {html_output}")
    print(f"  Report: {report_output}")
    print()

    print("Recommended Substack workflow:")
    print("  1. Open the *_substack.docx file in Microsoft Word.")
    print("  2. Select the article content.")
    print("  3. Copy.")
    print("  4. Paste into the Substack editor.")
    print()
    print("If Word copy/paste behaves oddly, open the HTML file in")
    print("Chrome/Edge and copy the rendered article from there.")
    print()

    if missing:
        print(
            "WARNING: Some source phrases were not found."
        )
        print(
            "Check the generated *_link_report.txt file."
        )
        print()


if __name__ == "__main__":
    main()
