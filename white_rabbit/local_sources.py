from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass
from pathlib import Path

import pymupdf
from bs4 import BeautifulSoup
from docx import Document


SUPPORTED = {".pdf", ".docx", ".md", ".markdown", ".txt", ".csv", ".json", ".html", ".htm"}


@dataclass
class LocalDocument:
    path: Path
    title: str
    text: str


@dataclass
class TextChunk:
    document: LocalDocument
    index: int
    text: str
    score: float = 0.0


def read_local_document(path: Path) -> LocalDocument:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED:
        raise ValueError(f"Unsupported local source: {path}")
    if suffix == ".pdf":
        doc = pymupdf.open(path)
        text = "\n".join(page.get_text("text") for page in doc)
    elif suffix == ".docx":
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
    elif suffix in {".html", ".htm"}:
        soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
        text = soup.get_text("\n", strip=True)
    elif suffix == ".csv":
        rows = []
        with path.open("r", encoding="utf-8-sig", errors="ignore", newline="") as f:
            for row in csv.reader(f):
                rows.append(" | ".join(row))
        text = "\n".join(rows)
    elif suffix == ".json":
        obj = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
        text = json.dumps(obj, ensure_ascii=False, indent=2)
    else:
        text = path.read_text(encoding="utf-8-sig", errors="ignore")
    return LocalDocument(path=path.resolve(), title=path.stem.replace("_", " "), text=text)


def discover_local_documents(folder: Path) -> list[LocalDocument]:
    folder = Path(folder)
    if not folder.exists():
        return []
    docs: list[LocalDocument] = []
    for path in sorted(folder.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED:
            try:
                docs.append(read_local_document(path))
            except Exception as exc:
                print(f"WARNING: could not read local source {path}: {exc}")
    return docs


def chunk_document(doc: LocalDocument, max_chars: int = 14000, overlap: int = 1200) -> list[TextChunk]:
    text = re.sub(r"\n{3,}", "\n\n", doc.text).strip()
    if not text:
        return []
    chunks: list[TextChunk] = []
    start = 0
    idx = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        if end < len(text):
            boundary = text.rfind("\n\n", start + max_chars // 2, end)
            if boundary > start:
                end = boundary
        chunks.append(TextChunk(document=doc, index=idx, text=text[start:end]))
        idx += 1
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return chunks


def _terms(text: str) -> set[str]:
    return {w.lower() for w in re.findall(r"[A-Za-z0-9][A-Za-z0-9_-]{2,}", text) if len(w) >= 4}


def rank_chunks(chunks: list[TextChunk], topic_and_queries: str, limit: int) -> list[TextChunk]:
    q = _terms(topic_and_queries)
    for c in chunks:
        t = _terms(c.text[:25000])
        overlap = len(q & t)
        # Give a modest boost to chunks with exact proper-noun-like topic tokens.
        c.score = overlap / max(1, len(q))
    return sorted(chunks, key=lambda c: c.score, reverse=True)[:limit]
